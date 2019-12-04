#  must be run with python 2.7
import docx
from bs4 import BeautifulSoup as bs
import requests
import datetime
from datetime import date
import time
import sys
import os
import colorama
from colorama import Fore, Back, Style
colorama.init()

# -*- coding: utf-8 -*-

word_filename = "ge_pricewatch.docx"
p_history_filename = "purchase_history" 
# interesting item url list, currently must be manually put in because of odd url names...
watchlist = {
# "iron bar": "http://services.runescape.com/m=itemdb_rs/Iron+bar/viewitem?obj=2351",
# "iron ingot": "http://services.runescape.com/m=itemdb_rs/Steel+ingot/viewitem?obj=47074",
# "iron platebody": "http://services.runescape.com/m=itemdb_rs/Iron+platebody/viewitem?obj=1115",
# "steel bar": "http://services.runescape.com/m=itemdb_rs/Steel+bar/viewitem?obj=2353", 
# "steel platebody": "http://services.runescape.com/m=itemdb_rs/Steel+platebody/viewitem?obj=45462", 
# "steel ingot": "http://services.runescape.com/m=itemdb_rs/Steel+ingot/viewitem?obj=47074", 
"grapes": "http://services.runescape.com/m=itemdb_rs/Grapes/viewitem?obj=1987",
# "soft clay": "http://services.runescape.com/m=itemdb_rs/Soft+clay/viewitem?obj=1761", 
# "copper ore": "http://services.runescape.com/m=itemdb_rs/Copper+ore/viewitem?obj=436", 
# "iron ore": "http://services.runescape.com/m=itemdb_rs/Iron+ore/viewitem?obj=440", 
"spider silk": "http://services.runescape.com/m=itemdb_rs/Spider+silk/viewitem?obj=25547", 
# "yew incense sticks": "http://services.runescape.com/m=itemdb_rs/Yew+incense+sticks/viewitem?obj=47690", 
"raw chicken": "http://services.runescape.com/m=itemdb_rs/Raw+chicken/viewitem?obj=2138"
}

# will make a file, or overwrite one of the same name
def log_purchase_history(f_name, info_to_save):
# defaults to making/opening the file in cwd
  with open(f_name, 'w') as f:
    f.write(info_to_save)

# simple read function, returning all lines read
def rd_f_nreturn_lines(f_name):
  with open(f_name, 'r') as f:
    lines = f.readlines()
  return lines

# NOTE: reference of getting soup with "requests" module instead of urllib
def get_soup(url):
  headers = {"Connection":"keep-alive",'User-Agent': 'Mozilla/5.0'}
  html_page = requests.get(url, headers=headers)
  return bs(html_page.text, "html.parser")

def find_price_info(key, page_soup):
  try:
    ge_price = page_soup.find("div", class_="stats").h3.span.text
    return ge_price
  except:
    print("[ - ] !!! " + key + " text info not found")

# crawl to each page in the watchlist and get the price data
def crawl_and_extract_prices(watchlist):
  price_dict = {}
  for k in watchlist:
    item_url = watchlist[k]
    page_soup = get_soup(item_url)
    price_amount = find_price_info(k, page_soup)
    # get an int from the string price value, deleting the commas
    price_dict[k] = int(price_amount.replace(",",""))
  return price_dict

# gets the paragraph version of price changes
def find_changes(watchlist): 
  # Document date line-header for reference
  changelist = ("[[[ Scanned on: " + str(date.today()) + "]]]\n")
  for key in watchlist:
    page_soup = get_soup(watchlist[key])
    curr_price = find_price_info(key, page_soup)
    changes_data = page_soup.find("div", class_="stats").ul.find_all("li")
    # colorized terminal output
    print((watchlist[key] + "\n"  + key + " : " + Fore.GREEN + curr_price +  "gp" + Style.RESET_ALL + "\n--- --- ---"))
    # non-color version added to paragraph
    changelist += (watchlist[key] + "\n"  + key + " : " + curr_price  + "gp" + "\n--- --- ---")
    # idx 2-3 are 3month and 6month histories in changes_data
    for line in changes_data[0:1]:
      txt = line.text
      changelist += ("\n " + txt)
    changelist += "--- --- --- \n"
  return changelist
 
# NOTE: This prints to an existing file, a new file is docx.Document()  [with no args]
# This adds more and more lines to the document, keeping a history of each scan
def print_to_word(txt_block, file1):
  doc = docx.Document(file1)
  doc.add_paragraph(txt_block)
  doc.save(file1)

# PRICE CHECKER FUNCTION NOT IMPLEMENTED
# def find_current_price(current_d):
#   target = raw_input("Checking price for what item?: ")
#   if target.lower() != "x":
#     purchased_name = target
#     try:
#       # ~~~~~~~~~~~
#       print(target + " currently has a price of " current_d[purchased_name] + "gp")
#       # run a function here to move that purchase to the purchase holder, which should account for old buys
#     except: 
#       print(target + " ... no price found")

# interactivity loop NOT IMPLEMENTED YET
def ask_about_purchases(current_d):
  purchased_name = ""
  purchased_quantity = 0
  purchased_price = 0
  individual_price = 0
  bought_item = raw_input("What did you buy? [x] for next: ")
  if bought_item.lower() != "x":
    purchased_name = bought_item
    try:
      print(bought_item + " at " + str(current_d[bought_item]) + "gp each ... ... ...")
      individual_price += current_d[bought_item]
      # can loop here for user input selection
      try:
        quantity = int(raw_input("how many units of " + bought_item + " did you buy? "))
        if quantity < 1:
          "None bought, cancelling recording a purchase"
        else:
          purchased_price += (quantity * current_d[bought_item])
          purchased_quantity += quantity
          print(str(quantity) + " units of " + bought_item + " bought at " + str(current_d[bought_item]) + "gp ea. totaling to " + str(purchased_price) + "gp")
      except:
        "invalid number entered"
    except: 
      print(bought_item + " not found in item list")
      # return purchase data with a timestamp
  return [purchased_name, individual_price, purchased_quantity, purchased_price, datetime.datetime.now().strftime("%Y-%m-%d %H:%M")]

def add_purchase_info_to_temp(tempfile1, buy_info_list):
  with open(tempfile1, 'w') as f:
    for i in buy_info_list:
      f.write(str(i).strip() + "\n")
# ///////////

  # def find_expenditure(i, q, prices_dict):
  # based on an input item, i, and quantity, q, determine the cost of the items

  # /////////////////

print("crawling to get prices...")
curr_prices_dict = crawl_and_extract_prices(watchlist)

price_paragraph = find_changes(watchlist)
# print_to_word(price_paragraph, word_filename)

buy_info = ask_about_purchases(curr_prices_dict)
# read the item history from temp file
purchase_history = rd_f_nreturn_lines(p_history_filename)
print(purchase_history)
# make sure there was an actual purchase, or else, skip this step, protects against mistypings and false entries
if buy_info[1] != 0:
  # update purchase history with new stuff
  for x in buy_info:
    purchase_history.append(x)
  print("/// updating purchases")
  print(purchase_history)
  add_purchase_info_to_temp(p_history_filename, purchase_history)

# //////////////////////
# print("~~~ expenditures updated")
# add_purchase_info_to_temp(purchase_history, temp_info_f_name)
# # get a list formed from the
# updated_purchases = rd_f_nreturn_lines(temp_info_f_name)
# print(updated_purchases)
# print("----")
# # can get an int version of any item figure how to iterate through the total list, then find index of that name ot access other info
# # return historical price
# # compare historical price to current price at same quantity
# calculate differences

# DELETE FUNCTION VIA NAME AND ALL THE IDX TRAILING IT

# SAMPLE ACCESS
# print(int(updated_purchases[2]))


# write_purchases_to_file(curr_prices_dict)
