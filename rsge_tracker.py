#  must be run with python 2.7
import docx
from bs4 import BeautifulSoup as bs
import requests
from datetime import date
import sys
import colorama
from colorama import Fore, Back, Style
colorama.init()

# -*- coding: utf-8 -*-

# interesting item url list

watchlist = {
"iron bar": "http://services.runescape.com/m=itemdb_rs/Iron+bar/viewitem?obj=2351", 
"iron ingot": "http://services.runescape.com/m=itemdb_rs/Steel+ingot/viewitem?obj=47074",
"iron platebody": "http://services.runescape.com/m=itemdb_rs/Iron+platebody/viewitem?obj=1115"
# "steel bar": "http://services.runescape.com/m=itemdb_rs/Steel+bar/viewitem?obj=2353", 
# "steel platebody": "http://services.runescape.com/m=itemdb_rs/Steel+platebody/viewitem?obj=45462", 
# "steel ingot": "http://services.runescape.com/m=itemdb_rs/Steel+ingot/viewitem?obj=47074", 
# "grapes": "http://services.runescape.com/m=itemdb_rs/Grapes/viewitem?obj=1987", 
# "soft clay": "http://services.runescape.com/m=itemdb_rs/Soft+clay/viewitem?obj=1761", 
# "copper ore": "http://services.runescape.com/m=itemdb_rs/Copper+ore/viewitem?obj=436", 
# "iron ore": "http://services.runescape.com/m=itemdb_rs/Iron+ore/viewitem?obj=440", 
# "spider silk": "http://services.runescape.com/m=itemdb_rs/Spider+silk/viewitem?obj=25547", 
# "yew incense sticks": "http://services.runescape.com/m=itemdb_rs/Yew+incense+sticks/viewitem?obj=47690", 
# "raw chicken": "http://services.runescape.com/m=itemdb_rs/Raw+chicken/viewitem?obj=2138"
}
# init a placeholder to hold item names and prices scraped when run

# uClient = uReq(my_url)
# # get html content into a variable, then close the request out
# page_soup = bs(uClient.read(), "html.parser")
# uClient.close()

# NOTE: reference of getting soup with "requests" module instead of urllib
def get_soup(url):
  headers = {"Connection":"keep-alive",'User-Agent': 'Mozilla/5.0'}
  html_page = requests.get(url, headers=headers)
  return bs(html_page.text, "html.parser")

def find_price_info(key, page_soup):
  try:
    ge_price = page_soup.find("div", class_="stats").h3.span.text
    return ge_price
  except:
    print("[ - ] !!! " + key + " text info not found")

# crawl to each page in the watchlist and get the price data
def crawl_and_extract_prices(watchlist):
  price_dict = {}
  for k in watchlist:
    item_url = watchlist[k]
    page_soup = get_soup(item_url)
    price_amount = find_price_info(k, page_soup)
    # get an int from the string price value, deleting the commas
    price_dict[k] = int(price_amount.replace(",",""))
  return price_dict

def find_changes(watchlist): 
  # Document date scanned for reference
  changelist = ("[[[ Scanned on: " + str(date.today()) + "]]]\n")
  for key in watchlist:
    page_soup = get_soup(watchlist[key])
    curr_price = find_price_info(key, page_soup)
    changes_data = page_soup.find("div", class_="stats").ul.find_all("li")
    # colorized terminal output
    print((watchlist[key] + "\n"  + key + " : " + Fore.GREEN + curr_price +  "gp" + Style.RESET_ALL + "\n--- --- ---"))
    # non-color version added to paragraph
    changelist += (watchlist[key] + "\n"  + key + " : " + curr_price  + "gp" + "\n--- --- ---")
    # idx 2-3 are 3month and 6month histories in changes_data
    for line in changes_data[0:1]:
      txt = line.text
      changelist += ("\n " + txt)
    changelist += "--- --- --- \n"
  return changelist
 
# NOTE: This prints to an existing file, a new file is docx.Document()  [with no args]
def print_to_word(txt_block, file1):
  doc = docx.Document(file1)
  doc.add_paragraph(txt_block)
  doc.save(file1)

curr_prices_dict = crawl_and_extract_prices(watchlist)

price_paragraph = find_changes(watchlist)
# Below will show the daily up/down changes for the watchlist
# print(price_paragraph)
print_to_word(price_paragraph, "ge_pricewatch.docx")
