import sys
from bs4 import BeautifulSoup as bs
import requests
import time
import docx
import ftfy

# idx_url is the hub page linking to all units of interest
idx_url = "https://kancolle.fandom.com/wiki/Ship"
baseurl = "https://kancolle.fandom.com"
# test url for single page
# url = "https://kancolle.fandom.com/wiki/Prinz_Eugen"

def get_soup(url):
  headers = {"Connection":"keep-alive",'User-Agent': 'Mozilla/5.0'}
  html_page = requests.get(url, headers=headers)
  return bs(html_page.text, "html.parser")

# def get_tables(url):
#   p_soup = get_soup(url)
#   all_tables = p_soup.find_all("table", class_="wikitable")
#   tar_tables = all_tables[0:3]
#   return tar_tables

def get_url_and_name_arrays(idx_url):
  names_arr = []
  urls_arr = []
  idx_soup = get_soup(idx_url)
  # all_adv_tooltip_links = idx_soup.find_all("span")
  all_adv_tooltip_links = idx_soup.find_all("span", class_="advanced-tooltip")
  for x in all_adv_tooltip_links:
    names_arr.append(x.text)
  for x in all_adv_tooltip_links:
    urls_arr.append(x.find("a").get("href"))
  return [names_arr, urls_arr]

# check returns good-tags
def check_quotelength(p_soup):
  good_tags = []
  all_quotes = p_soup.find_all("tr", class_="shipquote")
  for tag in all_quotes:
    # ignore wasei-eigo lines where en and jp are combined with colspan 2
    if tag.find("td", colspan="2"):
      pass
      # print(tag.text.strip())
      # print("---")
    else:
      good_tags.append(tag)
  return good_tags

def get_headers(good_tags):
  header_arr = []
  # slice away "play " text with idx 5
  for tag in good_tags:
    header_text = (tag.find("td").text[5:].strip())
    header_arr.append("-----" + header_text + "-----")
  # print(header_arr[23])
  return header_arr

def get_auds(good_tags):
  aud_arr = []
  for tag in good_tags:
    aud_arr.append(tag.find("a").get("href"))
  # print(aud_arr[23])
  return aud_arr

def get_j_lines(good_tags):
  jpn_arr = []
  for tag in good_tags:
    j_line = tag.find("td", class_="shipquote-ja").text.strip()
    # mojibake purification and text correction for japanese letter encoding
    purified_line = ftfy.fix_encoding(j_line)
    fixed_j_line = ftfy.fix_text(purified_line)
    jpn_arr.append(fixed_j_line)
  # print(jpn_arr[28].encode("ascii", "replace"))
  return jpn_arr

def get_en_lines(good_tags):
  en_arr = []
  for tag in good_tags:
    has_eng = tag.find("td", class_="shipquote-en")
    if has_eng:
      en_line = has_eng.text.strip()
      # have to purify foreign chars even in en lines...
      purified_line = ftfy.fix_encoding(en_line)
      fixed_en_line = ftfy.fix_text(purified_line)
      en_arr.append(fixed_en_line)
    else:
      en_arr.append("No EN available, check page")
  # print(en_arr[28])
  return en_arr

def get_and_write_data(good_tags, doc):
  heads = get_headers(good_tags)
  auds = get_auds(good_tags)
  j_lines = get_j_lines(good_tags)
  en_lines = get_en_lines(good_tags)
  # iterate over arrs and LATER write to doc
  # for x in range(0, len(auds)):
  for x in range(0, len(auds)):
    doc.add_paragraph(heads[x])
    # print(heads[x])
    doc.add_paragraph(auds[x])
    # print(auds[x])
    doc.add_paragraph(j_lines[x])
    # print(j_lines[x].encode('ascii', 'replace'))
    doc.add_paragraph(en_lines[x])
    # print(en_lines[x])
  print("Finished writing unit report, moving on...")

def process_season_tag(tag):
# get header
  try:
    l1 = tag.find_all("a")
    header = l1[0].text
  except:
    header = "[-] no header"
# get audio
  try:
    l2 = tag.find("a", class_="internal")
    aud_link = l2.get("href")
  except:
    aud_link = "[-] no audio"
# get jp
  try:
    jp_data_tag = (tag.findNext("td"))
    jp_txt = jp_data_tag.text.strip()
    purified_line = ftfy.fix_encoding(jp_txt)
    fixed_j_line = ftfy.fix_text(purified_line)
  except:
    fixed_j_line = "[-] no Japanese"
# get en
  try:
    en1 = jp_data_tag.findNext("td")
    en_txt = en1.text.strip()
    purified_en_line = ftfy.fix_encoding(en_txt)
    fixed_en_line = ftfy.fix_text(purified_en_line)
  except:
    fixed_en_line = "[-] no English"
  return[header, aud_link, fixed_j_line, fixed_en_line]

def write_season_data(d_list, doc):
  doc.add_paragraph("-------" + d_list[0] + "---------------------")
  for datum in range(1, 4):
    doc.add_paragraph(d_list[datum])

def get_seasonal_quotes(p_soup, doc1):
  season_tags = []
  # If the season table is not a NoneType, scan info
  season_section = p_soup.find("span", id="Seasonal_Quotes")
  if season_section:
    season_table = p_soup.find("span", id="Seasonal_Quotes").parent.findNext("table")
    all_rows = season_table.find_all("tr")
    # get rid of initial header cells by slicing with 1 as the start point, eliminate last row irrelevant line
    for row in all_rows[1:-1]:
      if row.get("style") != "display:none":
        # placeholder for cell data within the rows...
        x1 = row.find_all("td")
        seasonal_data_list = process_season_tag(x1[0])
        write_season_data(seasonal_data_list, doc1)
  return season_tags
  
# ///////////////

def crawl_across_urls(baseurl, unit_names, unit_urls):
  for x in range(0, len(unit_urls)):
    curr_url = (baseurl + unit_urls[x])
    curr_name = unit_names[x]
    print("[+] " + curr_name + " ---- ")
    p_soup = get_soup(curr_url)
    # The check returns good tags
    good_tags = check_quotelength(p_soup)
    # personal preference for making filenames with underscore instead of space, slashes mess up pathing but are included in some names, meaning "or".
    fn = curr_name.replace(" ", "_").replace("/", "_or_")
    # make new document for each unit
    doc1 = docx.Document()
    get_and_write_data(good_tags, doc1)
    # !!! need to grab season quotes on a different targeting paradigm
    get_seasonal_quotes(p_soup, doc1)
    doc1.save(fn + ".docx")

# Process execution ///////////////
names_and_urls = get_url_and_name_arrays(idx_url)
unit_names = names_and_urls[0]
unit_urls = names_and_urls[1]
# The big function that runs it all -->
# !!! successfully grabbed up to 88
#  273 total ships
crawl_across_urls(baseurl, unit_names, unit_urls)
# This will take a while to crawl acrss the links
print("[+] All data successfully grabbed!")

# /// discover duplicate names among units...
# seen_names = {}
# multi_names = []
# for x in unit_names:
#     if x not in seen_names:
#       # assign unique names to a dictionary, give it value 1 as a counter
#       seen_names[x] = 1
#     else: 
#       if seen_names[x] == 1:
#         multi_names.append(x)
#       seen_names[x] += 1
      
# print(multi_names)

# ////////// testing sesonal targeting paradigm with single urls

# url = "https://kancolle.fandom.com/wiki/Warspite"
# url = "https://kancolle.fandom.com/wiki/Prinz_Eugen"
# url ="https://kancolle.fandom.com/wiki/Akagi"
# p_soup = get_soup(url)

# fn = "Akagi"
# good_tags = check_quotelength(p_soup)
# doc1 = docx.Document()
# get_and_write_data(good_tags, doc1)
# get_seasonal_quotes(p_soup, doc1)
# doc1.save(fn + ".docx")

  # /// sample tag data ////
  # <td> <a href="/wiki/Seasonal/Christmas_2015" title="Seasonal/Christmas 2015">Christmas 2015</a><br/><span class="audio-button"><a class="internal" href="https://vignette.wikia.nocookie.net/kancolle/images/3/37/Akagi_Christmas_2015.ogg/revision/latest?cb=20151208065750" title="Akagi Christmas 2015.ogg">Play</a></span>\n</td>, <td><span lang="ja" style="font-family: sans-serif;">\u3053\u308c\u306f...! \u304a\u3044\u3057\u3063\uff01\u3053\u308c\u3082\u3001\u30af\u30ea\u30b9\u30de\u30b9...! \u3044\u3044\u3067\u3059\u306d\u3002\u3042\u3063\u3001\u52a0\u8cc0\u3055\u3093\u3082\u3001\u98df\u3079\u3066\u307e\u3059\uff1f</span>\n</td>, <td>This is... Delicious. These are Christmas indeed. Ah, would you like some, Kaga-san?\n</td>, <td>\n</td>

#  //// sample tag from goodtags //////////

# <tr class="shipquote" data-shipquote-form="Base" data-shipquote-id="25">
# <td style="white-space: nowrap;"><span class="audio-button click-parent"><a class="internal" href="https://vignette.wikia.nocookie.net/kancolle/images/9/95/Iowa-Library.ogg/revision/latest?cb=20160503001038" title="Iowa-Library.ogg">Play</a></span> Library
# </td><td class="shipquote-ja" colspan="1" style=""><span lang="ja" style="font-family: sans-serif;">[[[Japanese gibberish here]]]</span>
# </td><td class="shipquote-en" style="">Hi! This is Iowa, name ship of the <i>Iowa</i>-class battleships. Although I am a high-speed battleship, I am heavily armed, and could be said to be the penultimate form of battleships. As the last of the battleship classes born in the USA, I will work hard in this fleet. Nice to meet you!
# </td><td>
# </td></tr>

# url = "https://kancolle.fandom.com/wiki/Shimushu"
# url = "https://kancolle.fandom.com/wiki/Nagato"
# url = "https://kancolle.fandom.com/wiki/Warspite"
# url = "https://kancolle.fandom.com/wiki/Prinz_Eugen"
# tar_tables = get_tables(url)

# p_soup = get_soup(url)
# data_arr1 = get_all_info(p_soup)
# # write all data to doc here!!! ---------
# for x in data_arr1:
#   print(x[3].encode("ascii", "replace"))

# unit_file_name = "t1"
# doc1 = docx.Document()
# process_target_tables(tar_tables, doc1)
# doc1.save(unit_file_name + ".docx")

# //////////////
# Preserved snippets of getting all english from soup instead of filtered tags:
# ///////////

# def get_auds_and_headers(p_soup):
#   all_quotes = p_soup.find_all("tr", class_="shipquote")
#   aud_arr = []
#   for tag in all_quotes:
#     aud_arr.append(tag.find("a").get("href"))
#   header_arr = []
#   for tag2 in all_quotes:
#     # slice away "play " text with idx 5
#     header_text = (tag2.find("td").text[5:].strip())
#     header_arr.append("-----" + header_text + "-----")
#   return [header_arr, aud_arr]

# def get_j_lines(p_soup):
#   jpn_arr = []
#   all_j = p_soup.find_all("td", class_="shipquote-ja")
#   for j_tag in all_j:
#     jpn_arr.append(j_tag.text.strip())
#   return jpn_arr

# def get_en_lines(p_soup):
#   en_arr = []
#   all_en = p_soup.find_all("td", class_="shipquote-en")
#   for en_tag in all_en:
#     en_arr.append(en_tag.text.strip())
#   return en_arr

# def get_all_info(p_soup):
#   temp = get_auds_and_headers(p_soup)
#   headers = temp[0]
#   auds = temp[1]
#   j_lines = get_j_lines(p_soup)
#   en_lines = get_en_lines(p_soup)
#   return [headers, auds, j_lines, en_lines]

# def crawl_across_urls(baseurl, unit_names, unit_urls):
#   for x in range(0, len(unit_urls)):
#     curr_url = (baseurl + unit_urls[x])
#     curr_name = unit_names[x]
#     print("[+] " + curr_name + " ---- ")
#     p_soup = get_soup(curr_url)
#     data_arr = get_all_info(p_soup)
#     aligned_lines = zip(*data_arr)
#     # personal preference for making filenames with underscore instead of space
#     fn = curr_name.replace(" ", "_").replace("/", "_or_")
#     # make new document for each unit
#     doc1 = docx.Document()
#     for n in range(0, len(data_arr[1])):
#       for x in range(0, 4):
#         pure_line = ftfy.fix_encoding(aligned_lines[n][x])
#         fixed_line = ftfy.fix_text(pure_line)
#         doc1.add_paragraph(fixed_line)
#     doc1.save(fn + ".docx")

# turl = "https://kancolle.fandom.com/wiki/Iowa"
# p_soup = get_soup(turl)
# data_arr = get_all_info(p_soup)
# aligned_lines = zip(*data_arr)