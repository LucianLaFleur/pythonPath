import sys
from bs4 import BeautifulSoup as bs
import requests
import time
import docx
import ftfy

# idx_url is the hub page linking to all units of interest
idx_url = "https://kancolle.fandom.com/wiki/Ship"
baseurl = "https://kancolle.fandom.com"
# test url for single page
# url = "https://kancolle.fandom.com/wiki/Prinz_Eugen"

def get_soup(url):
  headers = {"Connection":"keep-alive",'User-Agent': 'Mozilla/5.0'}
  html_page = requests.get(url, headers=headers)
  return bs(html_page.text, "html.parser")

def get_tables(url):
  p_soup = get_soup(url)
  all_tables = p_soup.find_all("table", class_="wikitable")
  tar_tables = all_tables[0:3]
  return tar_tables

def get_url_and_name_arrays(idx_url):
  names_arr = []
  urls_arr = []
  idx_soup = get_soup(idx_url)
  # all_adv_tooltip_links = idx_soup.find_all("span")
  all_adv_tooltip_links = idx_soup.find_all("span", class_="advanced-tooltip")
  for x in all_adv_tooltip_links:
    names_arr.append(x.text)
  for x in all_adv_tooltip_links:
    urls_arr.append(x.find("a").get("href"))
  return [names_arr, urls_arr]

def process_row_info(rowdata, n):
  tar_info = []
  # slice out this annoying "play" that's in the inner text of the audio button
  try:
    if n == 2:
      tar_info.append("----- "+ rowdata[0].text.strip() +" -----")
    else:
      tar_info.append("----- "+ rowdata[0].text[5:].strip() +" -----")
  except:
    print("check unit rowdata\n[-][-][-]")
  try:
    # audlink = rowdata[0].find("a").get("href")
    audlink = rowdata[0].find_all("a", text="Play")[0].get("href")
  except:
    audlink = "[-] no audio"
  tar_info.append(audlink)
  try:
    tar_info.append(rowdata[1].text.strip())
    # get English text
    tar_info.append(rowdata[2].text.strip())
  except:
    print("[-][-][-]Skipped irregular entry\n[-][-][-]\n[-][-][-]")
    pass
  return tar_info

def process_target_tables(tar_tables, doc):
  # only the first 3 tables are of interest
  for x in range(0, 3):
    tablenum = x
    tr_list = tar_tables[x].find_all("tr")
    print("--- Row list length :" + str(len(tr_list)))
    # time.sleep(1)
    # tr_list[0] is just the column headers... can be ignored
    for y in range(1, len(tr_list)):
      curr_rowdata = tr_list[y].find_all("td")
      print(curr_rowdata)
      print("^^^^^^^^^")
      tar_info = process_row_info(curr_rowdata, tablenum)

  # exceptions for wonky seasonal table
  # if any of these trigger, it's an empty data thing to be ignored
      try:
        if (tar_info[1].strip()) == "":
          pass
          # print("Detected nothing at idx 1!")
        elif (tar_info[2].strip()) == "":
          pass
          # print("Detected nothing at idx 2!")
        elif (tar_info[3].strip()) == "":
          pass
          # print("Detected nothing at idx 3!")
        else:
          for z in tar_info:
            try:
              purified_line = ftfy.fix_encoding(z)
              fixed_line = ftfy.fix_text(purified_line)
              doc.add_paragraph(fixed_line)
              # print(x.encode("ascii", "replace"))
              # run printing operation here
            except:
              print("Check for line info processing error...")
              time.sleep(1)
      except:
        pass
  print("[+] Page processing complete! ---")

# Process execution ///////////////
names_and_urls = get_url_and_name_arrays(idx_url)
unit_names = names_and_urls[0]
unit_urls = names_and_urls[1]

# Iterate over each url to get tabledata and extract
# tar_tables = get_tables(url)
# print("number of tables = "+ str(len(tar_tables)))

# def crawl_across_urls(baseurl, unit_names, unit_urls):
#   for x in range(0, len(unit_urls)):
#     curr_url = (baseurl + unit_urls[x])
#     curr_name = unit_names[x]
#     print(curr_url)
#     time.sleep(2)
#     tar_tables = get_tables(curr_url)
#     # personal preference for making filenames with underscore instead of space
#     fn = curr_name.replace(" ", "_")
#     # make new document for each unit
#     doc1 = docx.Document()
#     process_target_tables(tar_tables, doc1)
#     doc1.save(fn + ".docx")

# tn = unit_names[0:4]
# tu = unit_urls[0:4]

# crawl_across_urls(baseurl, tn, tu)

#  //// sample ////
url = "https://kancolle.fandom.com/wiki/Shimushu"
# url = "https://kancolle.fandom.com/wiki/Prinz_Eugen"
# tar_tables = get_tables(url)

def get_shit(url):
  p_soup = get_soup(url)
  all_tags = p_soup.find_all("tr", class_="shipquote")
  print(all_tags[5])
  return all_tags
  # return all_aud_tags[5].find_next()

# get aud link
# get jpn
# get en
get_shit(url)

# unit_file_name = "t1"
# doc1 = docx.Document()
# process_target_tables(tar_tables, doc1)
# doc1.save(unit_file_name + ".docx")

# //////////////

# find ship name from their url page
# test_urls = ["https://kancolle.fandom.com/wiki/Prinz_Eugen",
# "https://kancolle.fandom.com/wiki/Nagato",
# "https://kancolle.fandom.com/wiki/Mutsu"]

# for x in test_urls:
#   n1, n2, n3 = x.partition('wiki/')
#   name = n3.replace("_", " ")
#   print(name + "\n---")