from __future__ import unicode_literals
import docx
import ftfy
import sys
import codecs
from bs4 import BeautifulSoup as bs
import requests
import os
import re
# get time to allow for sleep... 
import time 
# -*- coding: utf-8 -*-

#  possible unicode in py3 ?
# import win_unicode_console

#  Need to work from base page, then find src_url

# sentence scraper for tatoeba
# string var for pathing with tatoeba's base route
tat = "https://tatoeba.org/"

# incomplete: dynamically generate url within each call of the "nav" method
# insert sentences page url below
url = 'https://tatoeba.org/eng/audio/index/jpn?page=2'
url_partition = "?page="

ex_sent_src = 'https://tatoeba.org//eng/sentences/show/2715375'

en_arr = ["he", "she", "you", "the", "it", "is", "am", "my", "i'm", "i'll", "i'd", "me", "we", "they", "has","were", "let's", "did", "are", "do", "not", "how", "could", "can", "can't", "not", "don't",  "didn't", "of", "him", "if", "her", "where", "as", "had", "for", "that", "try", "let", "aren't", "yet", "have", "go", "at", "may", "get", "up", "got", "will", "gone", "would", "what", "please", "hurry", "looks", "seems", "think"]
# wh- like who/what/where, shouldn't be needed because of state-of-being
# i, a, on, to crossover to other languages
# to, minimal risk, high reward because of Eng, infinitive

# need to prohibit german and other language terms that will make a false positive
# specific term instead of global search
ban_arr = ["ich", "eine", "das", "ein", "ist", "bist", "du", "euch","war", "der", "wie", "dies", "sind", "mir", "meinen", "zur", "mit", "dem", "denn", "nicht", "immer", "zeit", "noch", "schon", "ausge", "bitte", "danke", "vom", "heute", "joku"]
# ban_arr = ["noch", "gehe", "nicht", "war der", "wie war", "dies", "suche", "sind", "zur", "immer", "komm", "denn", "kuten", "isch", "zeit", "eine", "schon", "ausge", "bitte", "subite", "dass", "irgen", "kurz", "helf", "schul", "schl", "danke", "hinw", "verbr", "vom", "beim", "wurd", "genau", "gesp", "yh", "kuu", "heute", "mit dem" "nett", "gesa"]

def get_html_soup(url):
  headers = {"Connection":"keep-alive",'User-Agent': 'Mozilla/5.0'}

  html_page = requests.get(url, headers=headers)

  soup = bs(html_page.text, "html.parser")
  return soup

def find_url_ending(url, divider):
  # extract current page to curr_page var
  notimportant1, notimportant2, curr_page = url.partition(divider)
  # message showing what page you're on
  if curr_page:
    print("--- On page " + curr_page + " ---")
    return curr_page
  else:
    print("--- On page 1 ---")
    return str(divider) + "1"

# replace urlending
# list1 = range(2, 14)
# for y in list1:
#   print(y)

def find_link_tags_by_class(soup):
  elements = soup.find_all("a", class_="mainSentence")
  return elements

def extract_hrefs_from_taglist(base_route, taglist):
  extracted_href_list = []
  for tag in taglist:
    link = tag.get("href")
    # use a function to navigate to link and get english data from it
    extracted_href_list.append(base_route + link)
  return extracted_href_list

def find_furi(soup):
  elements = soup.find_all("span", class_="markup")
  el_list = []
  for el in elements:
    el_list.append(el.text)
  return el_list

def find_audio(soup):
  found_audios = []
  elements = soup.find_all("a", class_="audioAvailable")
  for el in elements:
    found_audios.append(el.get("href"))
  return found_audios

# Checks for any words in english array (At top of code) in the target sentence
#  old version --- (slow, searches entire list linearly)
# def search_for_en(sent, en_arr, ban_arr):
#   found_switch = False
#   lowercase_src_words = sent.lower()
#   for word in en_arr:
#     # if a keyword is found in the sentence, turn switch to true
#     if word in lowercase_src_words:
#       found_switch = True
#       for x in ban_arr:
#         if x in lowercase_src_words:
#           return False
#   return found_switch

def search_for_en(sent, en_arr, ban_arr):
  sent_arr = sent.split()
  # i = 0
  # while i < len(sent_arr)
  for word in sent_arr:
    try: 
      l_word = word.lower()
      print("searching for : " + l_word)
      if l_word in en_arr:
        # if l_word in ban_arr:
        #   return False
        return True
      elif l_word in ban_arr:
        return False
    except:
      return False
    
# processing returns both the English and the sent src link
def process_src_page(link, en_arr, ban_arr):
  content_pg_soup = get_html_soup(link)
  # extract id-num from src, example id: 2715375
  # sentence_id = re.findall("[^\D]*$", link)[0]
  text_divs = content_pg_soup.find_all("div", class_="text")
  # jpn_line = (text_divs[2].text)
  #  Note ^^ above is extracted from url listing all audio sentences in Japanese, only reliable indexed sentence on src pg
  # Scans first 7 sentences under Japanese for an English one
  for n in range(3, 10):
    # !!!!!!!!!!!!!!!! CHANGED!!!
    # if search_for_en(text_divs[n].text, en_arr, ban_arr):
      if search_for_en(text_divs[n].text, en_arr, ban_arr):
        print("[ + ] Glorious! (extracting EN translation)")
        time.sleep(1)
        print("Moving to next target...")
        return text_divs[n].text.strip()
      else:
        print("--- eng not found at : " + str(n) + " ---")
        try:
          print(text_divs[n].text.strip())
        except:
          print("---skipped printing unprintable char---")
        print("Not EN, trying next entry ... \n")
        time.sleep(1)
      # return text_divs[n].text.strip() + "\n" + link + "\n"
      # str(n) <-- used to check the return index, made sure this iteration CAN loop through multiple places
  return "English line not found at : " + link 

# get all english sentences in a list, drawing from a list of links, and checking against inclusion and exclusion arrays
def get_eng_from_src(src_links, en_arr, ban_arr):
  collected_en = []
  for x in src_links:
    collected_en.append(process_src_page(x, en_arr, ban_arr))
  return collected_en 

# Incomplete: do the whole gig for each page from 2 to 12

pgnum = find_url_ending(url, url_partition)
print("pgnum is : " + pgnum)
html_soup = get_html_soup(url)
# 100 items collected per page, so each proper arr is 100 items long
# extract audio files from soup


# audio_list = find_audio(html_soup)
# furi_list = find_furi(html_soup)



# Extract all engish translations and put them in an array?
targeted_a_tags = find_link_tags_by_class(html_soup)
# again, "tat" is the base route for tatoeba https://tatoeba.org/
# Will get all links to sentence pages in an arr
src_links = extract_hrefs_from_taglist(tat, targeted_a_tags)
# slice src_links[0:13] to test first dozen, preventing data overflow.
translation_list = get_eng_from_src(src_links[0:22], en_arr, ban_arr)

print("---\n///results///\n---")

co1 = 0
for x in translation_list:
  print("--- sentence # " + str(co1) + " ---")
  try:
    print(x)
  except:
    print("false")
  co1 += 1

# # !!!!! write to external word document!!!
# # Note: this makes a newdoc, so be careful of the name
# doc = docx.Document()

# Any list could be used for length, as they should all be the same... 100 for main pages, and somewhat less for the last page
# for x in range(0, len(audio_list)):
#   purified_jpn = ftfy.fix_encoding(furi_list[x])
#   fixed_jpn = ftfy.fix_text(purified_jpn)
#   doc.add_paragraph(fixed_jpn)
#   doc.add_paragraph(translation_list[x])
#   doc.add_paragraph(audio_list[x] + "\n ---")

# # Note: Name of doc below
# doc.save("jpTest8.docx")

# print("Data extraction successful!")

# for tran in translation_list:
#   print(str(translation_list.index(tran)) + " --- " + tran)

# for spot in range(0, 10):
#   print(src_links[spot])

# ///////////////////
#  Extract English from a list of links
# ///////////////////
# make a test arr and add each returned item to test arr
# t1 = "https://tatoeba.org/eng/sentences/show/2232987"
# t2 = "https://tatoeba.org/eng/sentences/show/126740"
# t3 = "https://tatoeba.org/eng/sentences/show/218563"
# t4 = "https://tatoeba.org/eng/sentences/show/140994"
# t5 = "https://tatoeba.org/eng/sentences/show/4055090"
# t6 = "https://tatoeba.org/eng/sentences/show/3649764"
# t7 = "https://tatoeba.org/eng/sentences/show/4055090"
# t8 = "https://tatoeba.org/eng/sentences/show/3560578"
# t9 = "https://tatoeba.org/eng/sentences/show/3526163"

# t_arr = [t1, t2, t3, t4, t5, t6, t7, t8, t9]

# # Process links to test if eng or error will be printed out
# cou1 = 0
# collected_en = []

# for x in t_arr:
#   collected_en.append(process_src_page(x, en_arr, ban_arr))

# print(collected_en)

# //////////////
#  Operations on content pages

# FAILS documenting inaccessable info-bits ////////

#  FAILED Scan for all available links on page
# all_anchortags = content_pg_soup.find_all("a")
# counter2 = 0
# for anch in all_anchortags:
#   if anch.get("href"):
#     print("link at idx " + str(counter2) + ' ---  ' + anch.get("href"))
#     counter2 += 1

# FAILED , audio not available on scraped page...
# x9 = content_pg_soup.find_all("div", class_="audio")
# print(x9)

# FAILED: Iterating over all spans shows no japanese in content, doublt curly braces means shit's being called by a ruby on rails framework, inaccessable...
# allspans = content_pg_soup.find_all("span")
# counter1 = 0
# for span in allspans:
#   if span.text:
#     print("idx" + str(counter1) + " content: " + span.text)
#     counter1 += 1

# /////////////////
#  Writing to an external word doc
# /////////////////

# !!! Write from collected arrays, don't worry about iterating for this
# purify mojibake, fix text output
  # purified_jpn = ftfy.fix_encoding(jpn_line)
  # fixed_jpn = ftfy.fix_text(purified_jpn)

  # # Re-work to encapsulate doc creation then calling write_to_paper from within doc_creator

  # doc = docx.Document()
  # def write_to_paper(url, j, e):
  #   doc.add_paragraph(url)
  #   doc.add_paragraph(j)
  #   doc.add_paragraph(e)
  #   doc.add_paragraph("---- ")

  # write_to_paper(ex_sent_src, fixed_jpn, eng_line)

  # #  Document writing works!
  # # eng_line = "erngin"
  # # write_to_paper(ex_sent_src, fixed_jpn, eng_line)

  # doc.save('jpTest3.docx')
  # print("doc written successfully!")
  # # print(x1.decode("UTF-8"))


# # ///////////////////////////////////
