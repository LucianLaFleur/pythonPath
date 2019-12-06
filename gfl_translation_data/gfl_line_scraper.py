from __future__ import unicode_literals
import docx
import ftfy
import sys
from bs4 import BeautifulSoup as bs
import requests
import io
import os
import re
import time
# -*- coding: utf-8 -*-

library_page = "https://en.gfwiki.com/wiki/T-Doll_Index"
base_url = "https://en.gfwiki.com/wiki/"
src_root = "https://en.gfwiki.com/"

def get_soup(url):
  headers = {"Connection":"keep-alive",'User-Agent': 'Mozilla/5.0'}
  html_page = requests.get(url, headers=headers)
  return bs(html_page.text, "html.parser")

# finds the card data from the inxed
def find_initial_data_from_soup(soup):
  # target class carrying key info inside is "card-bg-small" on this site
  all_span_cards = soup.find_all("span", class_="card-bg-small")
  return all_span_cards

# now we start extracting specific data from the html containing stuff we're interested in
def get_names_from_data(data_list):
  name_list = []
  # Can shorten data here if you [sl:ice] the data list!!!
  for x in data_list:
    name = x.find("span", class_="name").text
    name_list.append(name)
  return name_list

def generate_urls_from_names(name_list, base_url):
  character_urls = [] 
  for name in name_list:
    named_url = (base_url + name)
    character_urls.append(named_url)
  return character_urls

def find_weapon_info(soup):
  try:
    wep_note = ""
    info1 = soup.find(text="Weapon Background").find_next(text="Weapon Background").find_next('p')
    wep_note += ("\t" + info1.text)
    cou = 0
    while True:
      next_tag_type = info1.next_element.name
      # end the loop when h2 is encountered
      if next_tag_type == "h2":
        # print("[+][+] --> h2 found! Ending design section search...")
        break
      # Will grab text from any p or ul elements
      elif next_tag_type == "p" or next_tag_type == "ul":
        txt_content = info1.next_element.text
        try:
          # get the text of the next element we're peeking at
          wep_note += ("\t" + txt_content)
        except:
          # if weird chars show up, just give me question marks
          txt_content = txt_content.encode('ascii', 'replace')
          wep_note += ("\t" + txt_content)
      elif cou >= 40:
        print("Weapon info loop exceeded 40 runs")
        break
      # else:
        # print("ignored non-target tag...")
      # make current element the next element
      info1 = info1.next_element
      # bump up saftey net counter
      cou += 1
  except:
    return "[-] No weapon details"
  return wep_note

# get all text content until you run into an h2 tag
def find_design_info(soup):
  try:
    #design_note is var for design info bits
    design_note = ""
    design_header = soup.find(id="Design")
    p1 = design_header.next_element.next_element.next_element
    design_note += ("\t" + p1.text)
    cou = 0
    while True:
      next_tag_type = p1.next_element.name
      # end the loop when h2 is encountered
      if next_tag_type == "h2":
        # print("[+][+] --> h2 found! Ending design section search...")
        break
      elif next_tag_type == "p" or next_tag_type == "ul":
        # print("----- text content: " )
        txt_content = p1.next_element.text
        try:
          # get the text of the next element we're peeking at
          design_note += ("\t" + txt_content)
        except:
          # if weird chars show up, just give me question marks
         txt_content = txt_content.encode('ascii', 'replace')
         design_note += ("\t" + txt_content)
      elif cou >= 48:
        print("Loop exceeded 48 runs")
        break
      # make current element the next element
      p1 = p1.next_element
      cou += 1
  except:
    return "[-] No design details"
  return design_note

# find all <li> 's since the entire data set is contained in them
def find_trivia_info(soup):
  try:
    # trivia_bullets is list of all trivia text bits
    trivia_bullets = soup.find(id="Trivia").find_next('ul').find_all('li', recursive=False)
    trivia_paragraph = ""
    # all_trivia_list = []
    for x in trivia_bullets:
      bare_text = x.text
      try:
        trivia_paragraph += ("\t" + bare_text)
      # all_trivia_list.append(bare_text)
      # print(bare_text + "\n")
      except:
        bare_text = ("\t" + bare_text.encode('ascii', 'replace'))
        trivia_paragraph += bare_text
    # return all_trivia_list
    return trivia_paragraph
  except:
    return "[-] No trivia details"

# Doll icon extaraction works
def get_doll_icon_src(data_list): 
  doll_icon_list = []
  for x in data_list:
    icon = x.find("img", class_="doll-image").get("src")
    doll_icon_list.append(icon)
  return doll_icon_list

# Search all info units on the soup for the char's homepage
def extract_char_info(soup):
  x = find_weapon_info(soup)
  y = find_design_info(soup)
  z = find_trivia_info(soup)
  raw_suite = ("---weapon design ---\n" + x + "\n --- Char design/attire --- \n" + y + "\n --- Trivia Bullets --- \n" + z)
  # NOTE: purifying mojibake??
  purified__suite = ftfy.fix_encoding(raw_suite)
  fixed_suite = ftfy.fix_text(purified__suite)
  return fixed_suite
  # returns a paragraph-long string


# NOTE: These don't necessarily have to be segregated to work, but the extra deliniaiton helps with clarity; presumed 6-tables are mod3
def check_quote_tables(all_tbody_tags):
  if len(all_tbody_tags) == 4:
    # Omits the [0] idx because it's irrelevant data
    # print("Regular quotes found")
    general_quotes = all_tbody_tags[1]
    combat_quotes = all_tbody_tags[2]
    event_quotes = all_tbody_tags[3]
    all_quotes = [general_quotes, combat_quotes, event_quotes]
    return all_quotes
  # Make exception for additional tables for MOD3
  elif len(all_tbody_tags) == 7:
    # print("Mod3 quotes found")
    general_quotes = all_tbody_tags[1]
    combat_quotes = all_tbody_tags[2]
    event_quotes = all_tbody_tags[3]
    mod3_general_quotes = all_tbody_tags[4]
    mod3_combat_quotes = all_tbody_tags[5]
    mod3_event_quotes = all_tbody_tags[6]
    all_quotes = [general_quotes, combat_quotes, event_quotes, mod3_general_quotes, mod3_combat_quotes, mod3_event_quotes]
    return all_quotes
  else:
    # just add all results into an arr and return it
    all_quotes = []
    # Also omits [0] idx because it's irrelevant data
    for x in range(1, len(all_tbody_tags)):
      all_quotes.append(all_tbody_tags[x])
    return all_quotes
    time.sleep(1)

# below grabs the audio src
def get_td_aud_src(rowdata):
  try:
    aud_href = rowdata.find("a", text="Play").get("href")
  # if aud_href:
    aud_src = src_root + aud_href
  # else:
  except:
    aud_src = "No audio available"
  return aud_src

# takes in a list of "td" from the soup, the row data, which gets sifted below
def get_info_from_rowdata(rd_list):
  try:
    # scan for length of list, review HTML src
    if len(rd_list) == 5:
    # idx 0 should be title, 2 -> Japanese, 4 -> English
    # UNUSED, not needed as title is in audiofile name ; checks for row title
    # if rd_list[0].text.strip():
    #   title_rd = rd_list[0].text.strip()
    # else:
    #   print("No title found")
    #   title_rd = "untitled row"
    # check for JPN content
      if rd_list[2].text.strip():
        jpn_rd = rd_list[2].text.strip()
        # slice out the unneeded "play" text that appears
        jpn_rd = jpn_rd[:-4]
      else:
        # print("No Jpn text found")
        jpn_rd = "[-] No Jpn trans"
      # check for English content
      if rd_list[4].text.strip():
        eng_rd = rd_list[4].text.strip()
      else:
        # print("No Eng text found")
        end_rd = "[-] No Eng trans"
      return [jpn_rd, eng_rd]
      # if the row doesn't include a leftmost cat, then its length is reduced to 4, targets shift down 1
      #  The "secretary" section makes this conditional necessary
    elif len(rd_list) == 4:
        if rd_list[1].text.strip():
          jpn_rd = rd_list[1].text.strip()
          jpn_rd = jpn_rd[:-4]
        else:
          jpn_rd = "[-] No Jpn trans"
        if rd_list[3].text.strip():
          eng_rd = rd_list[3].text.strip()
        else:
          end_rd = "[-] No Eng trans"
        return [jpn_rd, eng_rd]
    else:
      print("irregular table row entry length detected")
      time.sleep(3)
  except:
    # print("There appears to be no data there...")
    jpn_rd = "[-] No Jpn trans"
    eng_rd = "[-] No Eng trans"
    return [jpn_rd, eng_rd]

def write_rowdata_to_word(rowdata_list, doc):
  purified_j_line = ftfy.fix_encoding(rowdata_list[1])
  fixed_j_line = ftfy.fix_text(purified_j_line)
  # Japanese should be idx 1 of the rd_list
  doc.add_paragraph(rowdata_list[0])
  doc.add_paragraph(fixed_j_line)
  doc.add_paragraph(rowdata_list[2])
  doc.add_paragraph(rowdata_list[3])

def process_table_rows(rows, doc):
  # get rid of headers in first row
  del rows[0]
  # print("length: " + str(len(rows)) + " row entries detected")
# NOTE: [  :::  ] incomplete [  :::  ]
# ENCAPSULATE: for each item in the row table, extract row data and the audio
  for n in range(0, len(rows)):
    # get all the table-data from within the row
    rowdata = rows[n]
    # extract the audio link
    aud_src = get_td_aud_src(rowdata)
    rd_list = rowdata.find_all("td")
    # get JPN and EN translation
    t_lines = get_info_from_rowdata(rd_list)
      
    extracted_rd = [aud_src, t_lines[0], t_lines[1], ("Line: " + str(n) + " ------------------------------------------------")]
    write_rowdata_to_word(extracted_rd, doc)
    # show console what's being written to word document
    # for x in extracted_rd:
    #   print(x.encode('ascii', 'replace'))
    # print("~~~ idx [" + str(n) + "] done, scanning next entry ~~~)")
    # time.sleep(1)

# //// info extraction portion and method execution
tdoll_index_soup = get_soup(library_page)
data_list = find_initial_data_from_soup(tdoll_index_soup)
# print("[ + ] added " + str(len(data_list)) + " items to index data array...")
name_list = get_names_from_data(data_list)
icon_list = get_doll_icon_src(data_list)
char_pg_urls = generate_urls_from_names(name_list, base_url)

# //////// Test links //////////
# mod3 included
# t1 = "https://en.gfwiki.com/wiki/M4_SOPMOD_II"
# irregular, limited info, but normal num of tables, just some are absent content:
# t1 = "https://en.gfwiki.com/wiki/Jill"
# t1 = "https://en.gfwiki.com/wiki/C96"
# //////// END: Test links //////////

# iterate over all named units
# for x in range(0, len(name_list)):
for x in range(77, 79):
  curr_doll = name_list[x]
  curr_url = char_pg_urls[x]
  curr_icon = icon_list[x]
  print("---")
  print("scanning info for " + curr_doll.encode('ascii', 'replace') + "  --> " + curr_url.encode('ascii', 'replace'))
  time.sleep(1)
  character_soup = get_soup(curr_url)
  # print("grabbing character info, like trivia and background...")
  char_info_suite = extract_char_info(character_soup)
  quote_url = (curr_url + "/Quotes")
  character_quote_soup = get_soup(quote_url)

  all_tbody_tags = character_quote_soup.find_all("tbody")

  # make a new document for the unit
  tdoc = docx.Document()
  # give the name the doc will be saved under
  word_output_filename = curr_doll.encode('ascii', 'ignore') + "_translations.docx"
  # GETTING THE DUMB ICON IMAGE, OK?
  response = requests.get(curr_icon, stream=True)
  icon_img = io.BytesIO(response.content)
  tdoc.add_picture(icon_img)
  # write the document's header and the info suite at the top of the page
  tdoc.add_paragraph("Quotes for " + curr_doll + " : \n" + char_info_suite)
  # extract all the quote tables from the char page
  quote_tables = check_quote_tables(all_tbody_tags)
  try:
    print("Num of quote tables found: " + str((len(quote_tables))))
  except:
    print("Irregul")
  # quote_tables[0,1,2] => general, combat, event, may have extra if mod3 is active
  # Iterate over all detected tables
  for x in range(0, len(quote_tables)):
    gen_rows = quote_tables[x].find_all('tr')
    process_table_rows(gen_rows, tdoc)
  # save the document
  tdoc.save(word_output_filename)

print("---\n All scanning and printing successfully completed! ~ ~ ~")

# ///////////////////////////////////

# NOTE: purifying mojibake
# doc = docx.Document()
# for jpn in furi_list:
#   purified_jpn = ftfy.fix_encoding(jpn)
#   fixed_jpn = ftfy.fix_text(purified_jpn)
#   doc.add_paragraph(purified_jpn)
# doc.save('jpTest4.docx')

# Purify mojibake with a given line variable: jpn_line
# purified_jpn = ftfy.fix_encoding(jpn_line)
# fixed_jpn = ftfy.fix_text(purified_jpn)

# NOTE: .encode('utf-8') might be necessary if the doll's name has a weird char or backslash in their name...
#   buggytitle = atag.get('title')
#   encodedtitle = buggytitle.encode('utf-8')

# things we want in our output document:
#  T-doll name : tdoll_name
# Japanese
# English
# audio link
