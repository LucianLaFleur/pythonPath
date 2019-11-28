from __future__ import unicode_literals
import docx
import ftfy
import sys
import codecs
from bs4 import BeautifulSoup as bs
import requests
import os
import re
# get time to allow for sleep... 
import time 
# -*- coding: utf-8 -*-

#  filename to save doc as
f_name = "crawlJtest4.docx"

# sentence scraper for tatoeba
# string var for pathing with tatoeba's base route
href_root = "https://tatoeba.org/"
#  Note: some ID's are repeated, so... those should be skipped

# incomplete: dynamically generate url within each call of the "nav" method
# insert sentences page url below
base_url = 'https://tatoeba.org/eng/audio/index/jpn'
# insert the number for the last page (could be made a raw-input statement, bypassing console input by hard-coding here)
# Tested up to 3 pages working fine
last_page = 13
pages_list = range(2, (last_page + 1))
# url_partition only used to display current page being scanned in CLI, non-essential, along with find_url_ending()
url_partition = "?page="

# repeated 3573587

# somewhat prioritized by commonality, the lase several terms were added for oddballs found in testing
en_arr = ["he", "she", "you", "a", "the", "it", "my", "i'm", "i'll", "i'd", "we", "is", "they", "were", "let's", "did", "to", "are", "not", "could", "your", "his", "be", "her", "you're", "can", "can't", "what", "it's", "not", "has", "don't",  "didn't", "of", "him", "if", "her", "there", "where", "had", "for", "this", "too", "that", "any", "try", "let", "aren't", "yet", "have", "come", "go", "at", "should", "been", "may", "get", "up", "want", "got", "would", "by", "please", "please,", "please.", "isn't", "think", "how", "now", "looks", "looked", "their", "two", "must", "never", "made", "wish", "down", "down.", "went", "seldom", "you.", "put", "hasn't", "good", "call", "lost", "smile", "ask", "left", "how's", "said", "agree", "up.", "about", "i've", "some", "sleeping.", "sighed.", "why?", "swam.", "where's", "last.", "know.", "find", "everything.", "alone.", "forgive", "hug", "new", "until"]
# wh- like who/what/where, shouldn't be needed because of state-of-being
# i, an, on, me, will, was, to crossover to other languages
# to, minimal risk, high reward because of Eng, infinitive
# "do" minimal risk of omission, from central dusun?

# need to prohibit german and other language terms that will make a false positive
# specific term instead of global search
ban_arr = ["ich", "ihr", "er", "das", "eine", "sind", "wir","du", "ein", "ist", "bist", "euch", "der", "wie", "dies", "ihnen", "dir", "zur", "mit", "dem", "denn", "nicht", "noch", "nach", "bitte"]

def make_pages_url_list(base_page, partition, pgs_list):
  nav_urls = []
  for pg in pgs_list:
    numbered_url = base_page + partition + str(pg)
    nav_urls.append(numbered_url)
  return nav_urls
# for n in pages_list:
#   url = "https://tatoeba.org/eng/audio/index/jpn?page=" + str(n)
#   print(url + "\n ---")

def get_html_soup(url):
  headers = {"Connection":"keep-alive",'User-Agent': 'Mozilla/5.0'}

  html_page = requests.get(url, headers=headers)

  soup = bs(html_page.text, "html.parser")
  return soup

def find_url_ending(url, divider):
  # extract current page to curr_page var
  base_path, page_num_preamble, curr_page = url.partition(divider)
  # message showing what page you're on
  if curr_page:
    print("--- On page " + curr_page + " ---")
    return curr_page
  else:
    print("--- On page 1 ---")
    return "1"

def find_link_tags_by_class(soup):
  elements = soup.find_all("a", class_="mainSentence")
  return elements

def extract_hrefs_from_taglist(base_route, taglist):
  extracted_href_list = []
  for tag in taglist:
    link = tag.get("href")
    # use a function to navigate to link and get english data from it
    extracted_href_list.append(base_route + link)
  return extracted_href_list

def find_furi(soup):
  elements = soup.find_all("span", class_="markup")
  el_list = []
  for el in elements:
    el_list.append(el.text)
  return el_list

def find_audio(soup):
  found_audios = []
  elements = soup.find_all("a", class_="audioAvailable")
  for el in elements:
    found_audios.append(el.get("href"))
  return found_audios

# Checks for any words in english array (At top of code) in the target sentence
#  old version --- (slow, searches entire list linearly)
# def search_for_en(sent, en_arr, ban_arr):
#   found_switch = False
#   lowercase_src_words = sent.lower()
#   for word in en_arr:
#     # if a keyword is found in the sentence, turn switch to true
#     if word in lowercase_src_words:
#       found_switch = True
#       for x in ban_arr:
#         if x in lowercase_src_words:
#           return False
#   return found_switch

def search_for_en(sent, en_arr, ban_arr):
  sent_arr = sent.split()
  # i = 0
  # while i < len(sent_arr)
  for word in sent_arr:
    try: 
      l_word = word.lower()
      # print("searching for : " + l_word)
      if l_word in en_arr:
        if l_word in ban_arr:
          return False
        return True
      elif l_word in ban_arr:
        return False
    except:
      return False
    
# processing returns both the English and the sent src link
def process_src_page(link, en_arr, ban_arr):
  content_pg_soup = get_html_soup(link)
  # extract id-num from src, example id: 2715375
  # sentence_id = re.findall("[^\D]*$", link)[0]
  text_divs = content_pg_soup.find_all("div", class_="text")
  # jpn_line = (text_divs[2].text)
  #  Note ^^ above is extracted from url listing all audio sentences in Japanese, only reliable indexed sentence on src pg
  # Scans first 7 sentences under Japanese for an English one
  for n in range(3, 22):
    # !!!!!!!!!!!!!!!! CHANGED!!!
    # if search_for_en(text_divs[n].text, en_arr, ban_arr):
      if search_for_en(text_divs[n].text, en_arr, ban_arr):
        # print("[ + ] Glorious! (extracting EN translation)")
        # time.sleep(1)
        # print("Moving to next target...")
        return text_divs[n].text.strip()
      # else:
      #   try:
      #     print("--- No EN  : " + str(n) + " --> " + (text_divs[n].text.strip())
      #   except:
          # print("---skipped printing unprintable char---")
        # print("trying next entry ... \n")
        # time.sleep(1)
  # if all places in range checked, signal nothing was found :
  print("English not detected at " + link)
  return "#######\n English line not found at : " + link + " \n#########" 

# get all english sentences in a list, drawing from a list of links, and checking against inclusion and exclusion arrays
def get_eng_from_src(src_links, en_arr, ban_arr):
  collected_en = []
  for x in src_links:
    collected_en.append(process_src_page(x, en_arr, ban_arr))
  return collected_en 

def console_tran_res(translation_list):
  print("---\n///results///\n---")
  co1 = 0
  for x in translation_list:
    print("--- sentence # " + str(co1) + " ---")
    try:
      print(x)
    except:
      print("false")
    co1 += 1

def run_scanwrite_suite(doc, href_root, url, url_partition):
  pgnum = find_url_ending(url, url_partition)
  print("pgnum is : " + str(pgnum))
  print("Scanning soup")
  html_soup = get_html_soup(url)
  print("Building data lists")
  audio_list = find_audio(html_soup)
  furi_list = find_furi(html_soup)
  targeted_a_tags = find_link_tags_by_class(html_soup)
  src_links = extract_hrefs_from_taglist(href_root, targeted_a_tags)
  print("Extracting translations and checking for English")
  translation_list = get_eng_from_src(src_links, en_arr, ban_arr)
  # doc writing portion below, assumes creation and saving is external
  print("Writing data to word document")
  time.sleep(1)
  for x in range(0, len(audio_list)):
    purified_jpn = ftfy.fix_encoding(furi_list[x])
    fixed_jpn = ftfy.fix_text(purified_jpn)
    doc.add_paragraph(fixed_jpn)
    doc.add_paragraph(translation_list[x])
    doc.add_paragraph(audio_list[x] + "\n ---")

def extract_across_pages(href_root, base_url, url_partition, pages_list, f_name):
  doc = docx.Document()
  url_list = make_pages_url_list(base_url, url_partition, pages_list)
  # note url list excludes first page, needs to be called with --> base url
  run_scanwrite_suite(doc, href_root, base_url, url_partition)
  print("[ + ] Initial page finished")
  # then run scan on all listed items
  countup = 2
  for url in url_list:
    run_scanwrite_suite(doc, href_root, url, url_partition)
    print("[ + ] page " + str(countup) + " finished")
    countup += 1
  # after scanner loop runs, save document
  doc.save(f_name)
  # console_tran_res(translation_list)

extract_across_pages(href_root, base_url, url_partition, pages_list, f_name)

# /////////////////////

# pgnum = find_url_ending(url, url_partition)
# print("pgnum is : " + pgnum)
# html_soup = get_html_soup(url)
# # 100 items collected per page, so each proper arr is 100 items long
# # extract audio files from soup
# audio_list = find_audio(html_soup)
# furi_list = find_furi(html_soup)

# # Extract all engish translations and put them in an array?
# targeted_a_tags = find_link_tags_by_class(html_soup)
# # Will get all links to sentence pages in an arr
# src_links = extract_hrefs_from_taglist(href_root, targeted_a_tags)
# # slice src_links[0:28] to test first 28, preventing data overflow.
# translation_list = get_eng_from_src(src_links[0:28], en_arr, ban_arr)

# print("---\n///results///\n---")

# co1 = 0
# for x in translation_list:
#   print("--- sentence # " + str(co1) + " ---")
#   try:
#     print(x)
#   except:
#     print("false")
#   co1 += 1

# # !!!!! write to external word document!!!
# # Note: this makes a newdoc, so be careful of the name
# doc = docx.Document()
# Any list could be used for length, as they should all be the same... 100 for main pages, and somewhat less for the last page
# for x in range(0, len(audio_list)):
#   purified_jpn = ftfy.fix_encoding(furi_list[x])
#   fixed_jpn = ftfy.fix_text(purified_jpn)
#   doc.add_paragraph(fixed_jpn)
#   doc.add_paragraph(translation_list[x])
#   doc.add_paragraph(audio_list[x] + "\n ---")
# # Note: Name of doc below
# doc.save("jpCrawlTest1.docx")

print("Data execution successful!")

# ///////////////////////////////////
